"""
Script para procesar el diseño técnico GESTIC-833509_Diseno_Tecnico_Detalle_RD_v2.docx
Usa el modelo qwen2.5 de Ollama
"""

import os
import sys
import re
import math
from docx import Document
import pandas as pd
from openpyxl import load_workbook

# Configurar encoding para Windows
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

# ============================================================================
# CONFIGURACIÓN
# ============================================================================

ARCHIVO_WORD = r"C:\Users\Nitropc\Desktop\LANGCHAIN\GESTIC-833509_Diseno_Tecnico_Detalle_RD_v2.docx"
EXCEL_TEMPLATE = r"C:\Users\Nitropc\Desktop\LANGCHAIN\GESTIC-XXXXXX Proyecto - Descripción.xlsx"
ARCHIVO_SALIDA = "estimacion_gestic_833509_RD.xlsx"

# ============================================================================
# PASO 1: LEER EL DOCUMENTO WORD
# ============================================================================

print("="*80)
print("📄 PROCESANDO DISEÑO TÉCNICO GESTIC-833509")
print("="*80 + "\n")

print("🔍 Leyendo documento Word...")
doc = Document(ARCHIVO_WORD)

# Extraer todo el texto
texto_completo = []
for parrafo in doc.paragraphs:
    if parrafo.text.strip():
        texto_completo.append(parrafo.text)

# Extraer tablas
print(f"📊 Encontradas {len(doc.tables)} tablas en el documento\n")

tablas_info = []
for i, tabla in enumerate(doc.tables):
    print(f"   Tabla {i+1}:")
    filas_tabla = []
    for j, fila in enumerate(tabla.rows):
        celdas = [celda.text.strip() for celda in fila.cells]
        filas_tabla.append(celdas)
        if j < 3:  # Mostrar primeras 3 filas
            print(f"      {' | '.join(celdas[:5])}")  # Primeras 5 columnas
    tablas_info.append(filas_tabla)
    print()

texto_completo_str = "\n".join(texto_completo)

# Guardar el texto extraído para inspección
with open("texto_extraido.txt", "w", encoding="utf-8") as f:
    f.write(texto_completo_str)
    f.write("\n\n=== TABLAS ===\n\n")
    for i, tabla in enumerate(tablas_info):
        f.write(f"\n--- Tabla {i+1} ---\n")
        for fila in tabla:
            f.write(" | ".join(fila) + "\n")

print(f"✅ Texto extraído guardado en texto_extraido.txt\n")
print(f"📝 Total caracteres: {len(texto_completo_str)}")
print(f"📝 Total párrafos: {len(texto_completo)}\n")

# ============================================================================
# PASO 2: ANÁLISIS MANUAL DEL CONTENIDO
# ============================================================================

print("="*80)
print("🔍 ANALIZANDO COMPONENTES DEL DISEÑO TÉCNICO")
print("="*80 + "\n")

# Buscar patrones comunes en diseños técnicos
componentes_encontrados = {
    "modelos": [],
    "servicios": [],
    "componentes": [],
    "vistas": [],
    "apis": [],
    "tablas_bd": []
}

# Patrones para detectar componentes
texto_lower = texto_completo_str.lower()

# Buscar menciones de entidades/modelos
patrones_modelo = [
    r"(?:modelo|entidad|entity|clase|class)\s+[:\-]?\s*(\w+)",
    r"tabla\s+(\w+)",
    r"(?:^|\n)(\w+Entity)",
    r"(?:^|\n)(\w+Model)"
]

for patron in patrones_modelo:
    matches = re.finditer(patron, texto_completo_str, re.IGNORECASE | re.MULTILINE)
    for match in matches:
        nombre = match.group(1)
        if len(nombre) > 2 and nombre not in componentes_encontrados["modelos"]:
            componentes_encontrados["modelos"].append(nombre)

# Buscar servicios
patrones_servicio = [
    r"(?:servicio|service)\s+[:\-]?\s*(\w+)",
    r"(\w+Service)",
    r"(\w+Controller)"
]

for patron in patrones_servicio:
    matches = re.finditer(patron, texto_completo_str, re.IGNORECASE)
    for match in matches:
        nombre = match.group(1)
        if len(nombre) > 2 and nombre not in componentes_encontrados["servicios"]:
            componentes_encontrados["servicios"].append(nombre)

# Buscar componentes frontend
patrones_componente = [
    r"(?:componente|component)\s+[:\-]?\s*(\w+)",
    r"(\w+Component)",
    r"página\s+(\w+)",
    r"pantalla\s+(\w+)"
]

for patron in patrones_componente:
    matches = re.finditer(patron, texto_completo_str, re.IGNORECASE)
    for match in matches:
        nombre = match.group(1)
        if len(nombre) > 2 and nombre not in componentes_encontrados["componentes"]:
            componentes_encontrados["componentes"].append(nombre)

# Mostrar resumen
print("📦 COMPONENTES DETECTADOS:\n")
print(f"   🗄️  Modelos/Entidades: {len(componentes_encontrados['modelos'])}")
for modelo in componentes_encontrados['modelos'][:10]:
    print(f"      - {modelo}")
if len(componentes_encontrados['modelos']) > 10:
    print(f"      ... y {len(componentes_encontrados['modelos']) - 10} más")

print(f"\n   🔧 Servicios: {len(componentes_encontrados['servicios'])}")
for servicio in componentes_encontrados['servicios'][:10]:
    print(f"      - {servicio}")
if len(componentes_encontrados['servicios']) > 10:
    print(f"      ... y {len(componentes_encontrados['servicios']) - 10} más")

print(f"\n   🎨 Componentes: {len(componentes_encontrados['componentes'])}")
for comp in componentes_encontrados['componentes'][:10]:
    print(f"      - {comp}")
if len(componentes_encontrados['componentes']) > 10:
    print(f"      ... y {len(componentes_encontrados['componentes']) - 10} más")

# ============================================================================
# PASO 3: ESTIMACIÓN BASADA EN COMPLEJIDAD DETECTADA
# ============================================================================

print("\n" + "="*80)
print("🧮 CALCULANDO ESTIMACIÓN")
print("="*80 + "\n")

# Cargar ponderaciones
ponderaciones = pd.read_excel(EXCEL_TEMPLATE, sheet_name='Ponderaciones', header=0)

# Estimación conservadora por componente
class ComponenteEstimado:
    def __init__(self, tipo, nombre):
        self.tipo = tipo
        self.nombre = nombre
        self.propiedades = 0
        self.metodos = 0
        self.eventos = 0
        self.integraciones = 0
        self.reglas_negocio = 0
        self.horas = 0

componentes_estimados = []

# Modelos (valores conservadores)
for modelo in componentes_encontrados['modelos'][:20]:  # Limitar a 20 más relevantes
    comp = ComponenteEstimado("Modelo de datos", modelo)
    comp.propiedades = 8  # Promedio de propiedades
    comp.metodos = 2
    comp.reglas_negocio = 5
    # Calcular horas (según ponderaciones GESTIC) - REDONDEO AL SIGUIENTE ENTERO
    comp.horas = math.ceil((comp.propiedades * 10 + comp.metodos * 15 + comp.reglas_negocio * 60) / 60)
    componentes_estimados.append(comp)

# Servicios
for servicio in componentes_encontrados['servicios'][:15]:
    comp = ComponenteEstimado("Servicios", servicio)
    comp.propiedades = 3
    comp.metodos = 12
    comp.integraciones = 3
    # REDONDEO AL SIGUIENTE ENTERO
    comp.horas = math.ceil((comp.propiedades * 10 + comp.metodos * 15 + comp.integraciones * 45) / 60)
    componentes_estimados.append(comp)

# Componentes frontend
for componente in componentes_encontrados['componentes'][:15]:
    comp = ComponenteEstimado("Componentes (TS)", componente)
    comp.propiedades = 8
    comp.metodos = 15
    comp.eventos = 10
    comp.integraciones = 5
    comp.reglas_negocio = 12
    # REDONDEO AL SIGUIENTE ENTERO
    comp.horas = math.ceil((comp.propiedades * 10 + comp.metodos * 30 + comp.eventos * 30 +
                  comp.integraciones * 45 + comp.reglas_negocio * 60) / 60)
    componentes_estimados.append(comp)

# Calcular totales - TODOS CON REDONDEO AL SIGUIENTE ENTERO
total_horas_base = sum(c.horas for c in componentes_estimados)

# Aplicar factores (asumiendo Media incertidumbre, Sí acoplamiento, Mid seniority)
factor_seniority = 1.5  # Mid
horas_ajustadas = math.ceil(total_horas_base * factor_seniority)

factor_incertidumbre = 0.15  # Media
horas_incertidumbre = math.ceil(horas_ajustadas * factor_incertidumbre)

factor_acoplamiento = 0.15  # Sí
horas_acoplamiento = math.ceil(horas_ajustadas * factor_acoplamiento)

total_final = math.ceil(horas_ajustadas + horas_incertidumbre + horas_acoplamiento)

# Distribución por fases - TODAS REDONDEADAS AL SIGUIENTE ENTERO
horas_analisis = math.ceil(horas_ajustadas * 0.10)
horas_construccion = math.ceil(horas_ajustadas * 0.60)
horas_pruebas = math.ceil(horas_ajustadas * 0.20)
horas_documentacion = math.ceil(horas_ajustadas * 0.05)
horas_peer_review = math.ceil(horas_ajustadas * 0.05)

print("📊 RESUMEN DE ESTIMACIÓN:\n")
print(f"   Componentes analizados: {len(componentes_estimados)}")
print(f"      - Modelos: {len([c for c in componentes_estimados if c.tipo == 'Modelo de datos'])}")
print(f"      - Servicios: {len([c for c in componentes_estimados if c.tipo == 'Servicios'])}")
print(f"      - Componentes TS: {len([c for c in componentes_estimados if c.tipo == 'Componentes (TS)'])}")
print(f"\n   Horas base: {total_horas_base}h")
print(f"   Horas ajustadas (Mid): {horas_ajustadas}h")
print(f"   + Incertidumbre (Media 15%): {horas_incertidumbre}h")
print(f"   + Acoplamiento (15%): {horas_acoplamiento}h")
print(f"   {'─'*40}")
print(f"   TOTAL FINAL: {total_final} horas")
print(f"\n   📅 Duración estimada:")
print(f"      - Con 1 desarrollador: {math.ceil(total_final/40)} semanas")
print(f"      - Con 2 desarrolladores: {math.ceil(total_final/80)} semanas")
print(f"      - Con 3 desarrolladores: {math.ceil(total_final/120)} semanas")
print(f"\n   📋 Distribución por fases:")
print(f"      - Análisis técnico: {horas_analisis}h")
print(f"      - Construcción: {horas_construccion}h")
print(f"      - Pruebas: {horas_pruebas}h")
print(f"      - Documentación: {horas_documentacion}h")
print(f"      - Peer Review: {horas_peer_review}h")

# ============================================================================
# PASO 4: EXPORTAR A EXCEL
# ============================================================================

print("\n" + "="*80)
print("📤 EXPORTANDO A EXCEL GESTIC")
print("="*80 + "\n")

try:
    # Cargar plantilla
    wb = load_workbook(EXCEL_TEMPLATE)

    # Escribir en Estimador_resumen
    ws_resumen = wb['Estimador_resumen']
    fila_inicio = 4

    ws_resumen[f'D{fila_inicio+1}'] = 0  # Gestión
    ws_resumen[f'D{fila_inicio+2}'] = horas_analisis
    ws_resumen[f'D{fila_inicio+3}'] = horas_construccion
    ws_resumen[f'D{fila_inicio+4}'] = horas_pruebas
    ws_resumen[f'D{fila_inicio+5}'] = horas_documentacion
    ws_resumen[f'D{fila_inicio+6}'] = horas_peer_review

    ws_resumen['C1'] = "Media"  # Incertidumbre
    ws_resumen['C2'] = "Sí"     # Acoplamiento

    # Escribir en Estimador_desglosado
    ws_desglosado = wb['Estimador_desglosado']
    fila_actual = 6

    for comp in componentes_estimados[:50]:  # Limitar a 50 componentes
        ws_desglosado[f'A{fila_actual}'] = comp.tipo
        ws_desglosado[f'B{fila_actual}'] = "SÍ"
        ws_desglosado[f'C{fila_actual}'] = comp.nombre
        ws_desglosado[f'D{fila_actual}'] = comp.propiedades
        ws_desglosado[f'E{fila_actual}'] = comp.metodos
        ws_desglosado[f'F{fila_actual}'] = comp.eventos
        ws_desglosado[f'G{fila_actual}'] = comp.integraciones
        ws_desglosado[f'H{fila_actual}'] = comp.reglas_negocio
        ws_desglosado[f'I{fila_actual}'] = comp.horas
        fila_actual += 1

    # Guardar
    wb.save(ARCHIVO_SALIDA)

    print(f"✅ Estimación exportada a: {ARCHIVO_SALIDA}")
    print(f"   Total componentes: {len(componentes_estimados)}")
    print(f"   Total horas: {total_final}h")

except Exception as e:
    print(f"❌ Error al exportar: {e}")

print("\n" + "="*80)
print("✅ PROCESO COMPLETADO")
print("="*80)
print(f"\nArchivos generados:")
print(f"   - texto_extraido.txt (análisis del documento)")
print(f"   - {ARCHIVO_SALIDA} (estimación GESTIC)")
print()
