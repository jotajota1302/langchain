"""
Agente de Estimación de Desarrollo con LangChain + Ollama
Lee diseños técnicos (PDF/Word) y genera estimaciones en plantilla Excel GESTIC
"""

import sys
import os
import re
import math
from typing import Dict, List, Any
from datetime import datetime

# Configurar codificación UTF-8 para Windows
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

# LangChain
from langchain_community.llms import Ollama
from langchain_classic.agents import AgentExecutor, create_react_agent
from langchain_core.tools import Tool
from langchain_core.prompts import PromptTemplate
from langchain_classic.memory import ConversationBufferMemory

# Procesamiento de documentos
import PyPDF2
from docx import Document
import pandas as pd
from openpyxl import load_workbook
from openpyxl.styles import Font, PatternFill, Alignment


# ============================================================================
# CONFIGURACIÓN
# ============================================================================

# Obtener el directorio raíz del proyecto (un nivel arriba de AGENTS/)
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# Rutas relativas usando la nueva estructura de carpetas
EXCEL_TEMPLATE_PATH = os.path.join(PROJECT_ROOT, "INPUT", "GESTIC-XXXXXX Proyecto - Descripción.xlsx")
PROMPT_ESTIMACION_PATH = os.path.join(PROJECT_ROOT, "PROMPTS", "agente-estimacion-desarrollo.md")

# Leer el prompt de estimación
with open(PROMPT_ESTIMACION_PATH, 'r', encoding='utf-8') as f:
    PROMPT_ESTIMACION = f.read()


# ============================================================================
# CLASES DE DATOS
# ============================================================================

class ComponenteEstimado:
    """Representa un componente con su estimación"""
    def __init__(self, tipo: str, nombre: str, descripcion: str = ""):
        self.tipo = tipo  # Modelo, Módulo, Servicio, Componente (TS), Vista HTML, Estilo SCSS
        self.nombre = nombre
        self.descripcion = descripcion
        self.es_nuevo = True  # True=Nuevo, False=Actualización
        self.propiedades = 0
        self.metodos = 0
        self.eventos = 0
        self.integraciones = 0
        self.reglas_negocio = 0
        self.horas_calculadas = 0
        self.complejidad = "Media"  # Baja, Media, Alta

    def calcular_horas(self, ponderaciones: pd.DataFrame) -> float:
        """Calcula las horas según ponderaciones"""
        # Obtener ponderaciones por tipo
        pond_tipo = ponderaciones[ponderaciones['MIS TIEMPOS'] == self.tipo]
        if pond_tipo.empty:
            return 0

        horas = 0
        horas += self.propiedades * float(pond_tipo['Propiedad'].values[0] or 0)
        horas += self.metodos * float(pond_tipo['Método'].values[0] or 0)
        horas += self.eventos * float(pond_tipo['Evento'].values[0] or 0)
        horas += self.integraciones * float(pond_tipo['Integración'].values[0] or 0)
        horas += self.reglas_negocio * float(pond_tipo['Regla \nNegocio'].values[0] or 0)

        # Convertir minutos a horas
        horas = horas / 60

        # Factor por nuevo/actualización
        factor_nuevo = 1.0 if self.es_nuevo else 1.5
        horas *= factor_nuevo

        # Redondear al siguiente entero
        self.horas_calculadas = math.ceil(horas)
        return self.horas_calculadas


class EstimacionProyecto:
    """Contiene toda la estimación del proyecto"""
    def __init__(self, nombre_proyecto: str):
        self.nombre = nombre_proyecto
        self.componentes: List[ComponenteEstimado] = []
        self.incertidumbre = "Baja"  # Nula, Baja, Media, Alta
        self.acoplamiento = False
        self.seniority = "Mid"  # Senior, Mid, Junior
        self.horas_gestion = 0
        self.horas_analisis = 0
        self.horas_construccion = 0
        self.horas_pruebas = 0
        self.horas_documentacion = 0
        self.horas_peer_review = 0

    def agregar_componente(self, componente: ComponenteEstimado):
        self.componentes.append(componente)

    def calcular_total_horas(self) -> float:
        return sum(c.horas_calculadas for c in self.componentes)

    def aplicar_factores_ajuste(self, ponderaciones: pd.DataFrame) -> Dict[str, float]:
        """Aplica factores de incertidumbre, acoplamiento y seniority"""
        total_base = self.calcular_total_horas()

        # Factor incertidumbre
        factores_incert = {
            "Nula": 0,
            "Baja": 0.07,
            "Media": 0.15,
            "Alta": 0.3
        }
        factor_incert = factores_incert.get(self.incertidumbre, 0.15)

        # Factor acoplamiento
        factor_acop = 0.15 if self.acoplamiento else 0

        # Factor seniority
        factores_sen = {
            "Senior": 1.0,
            "Mid": 1.5,
            "Junior": 2.5
        }
        factor_sen = factores_sen.get(self.seniority, 1.5)

        # Calcular totales - REDONDEO AL SIGUIENTE ENTERO
        horas_ajustadas = math.ceil(total_base * factor_sen)
        horas_incertidumbre = math.ceil(horas_ajustadas * factor_incert)
        horas_acoplamiento = math.ceil(horas_ajustadas * factor_acop)

        total_final = math.ceil(horas_ajustadas + horas_incertidumbre + horas_acoplamiento)

        return {
            "base": total_base,
            "ajustada_seniority": horas_ajustadas,
            "incertidumbre": horas_incertidumbre,
            "acoplamiento": horas_acoplamiento,
            "total_final": total_final
        }


# ============================================================================
# HERRAMIENTAS DEL AGENTE
# ============================================================================

# Variable global para almacenar la estimación actual
estimacion_actual = None


def leer_pdf(ruta_archivo: str) -> str:
    """
    Lee un archivo PDF y extrae todo el texto.
    Args:
        ruta_archivo: Ruta completa al archivo PDF
    Returns:
        Texto extraído del PDF
    """
    try:
        if not os.path.exists(ruta_archivo):
            return f"Error: El archivo {ruta_archivo} no existe"

        texto_completo = []

        with open(ruta_archivo, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_paginas = len(pdf_reader.pages)

            for num_pagina in range(num_paginas):
                pagina = pdf_reader.pages[num_pagina]
                texto = pagina.extract_text()
                texto_completo.append(f"--- Página {num_pagina + 1} ---\n{texto}")

        resultado = "\n\n".join(texto_completo)
        return f"PDF leído correctamente ({num_paginas} páginas):\n\n{resultado[:2000]}..."  # Primeros 2000 chars

    except Exception as e:
        return f"Error al leer PDF: {str(e)}"


def leer_word(ruta_archivo: str) -> str:
    """
    Lee un archivo Word (.docx) y extrae todo el texto.
    Args:
        ruta_archivo: Ruta completa al archivo Word
    Returns:
        Texto extraído del documento
    """
    try:
        if not os.path.exists(ruta_archivo):
            return f"Error: El archivo {ruta_archivo} no existe"

        doc = Document(ruta_archivo)

        texto_completo = []
        for i, parrafo in enumerate(doc.paragraphs):
            if parrafo.text.strip():
                texto_completo.append(parrafo.text)

        # Extraer tablas si las hay
        if doc.tables:
            texto_completo.append("\n--- TABLAS ENCONTRADAS ---")
            for i, tabla in enumerate(doc.tables):
                texto_completo.append(f"\nTabla {i+1}:")
                for fila in tabla.rows:
                    celdas = [celda.text for celda in fila.cells]
                    texto_completo.append(" | ".join(celdas))

        resultado = "\n".join(texto_completo)
        return f"Word leído correctamente:\n\n{resultado[:2000]}..."  # Primeros 2000 chars

    except Exception as e:
        return f"Error al leer Word: {str(e)}"


def extraer_componentes_tecnicos(texto_diseño: str) -> str:
    """
    Analiza el diseño técnico y extrae componentes (modelos, servicios, componentes, etc.)
    Args:
        texto_diseño: Texto del diseño técnico
    Returns:
        Análisis estructurado de componentes
    """
    global estimacion_actual

    # Inicializar nueva estimación
    estimacion_actual = EstimacionProyecto("Proyecto Estimado")

    # Patrones para detectar componentes
    patrones = {
        "Modelo de datos": r"(?:modelo|entidad|entity|tabla|table)s?\s+(\w+)",
        "Módulos": r"(?:módulo|module|paquete|package)s?\s+(\w+)",
        "Servicios": r"(?:servicio|service|api)s?\s+(\w+)",
        "Componentes (TS)": r"(?:componente|component)s?\s+(\w+)",
        "Vista HTML": r"(?:vista|view|página|page|pantalla|screen)s?\s+(\w+)",
        "Estilos SCSS": r"(?:estilo|style|tema|theme)s?\s+(\w+)"
    }

    componentes_encontrados = []

    for tipo, patron in patrones.items():
        matches = re.finditer(patron, texto_diseño, re.IGNORECASE)
        for match in matches:
            nombre = match.group(1)
            comp = ComponenteEstimado(tipo, nombre)
            # Valores por defecto estimados
            comp.propiedades = 3
            comp.metodos = 2
            comp.eventos = 1
            comp.integraciones = 1
            comp.reglas_negocio = 1
            estimacion_actual.agregar_componente(comp)
            componentes_encontrados.append(f"- {tipo}: {nombre}")

    if componentes_encontrados:
        return f"Componentes detectados:\n" + "\n".join(componentes_encontrados[:20])
    else:
        return "No se detectaron componentes automáticamente. Describe los componentes manualmente."


def agregar_componente_manual(especificacion: str) -> str:
    """
    Agrega un componente manualmente con su estimación.
    Formato: "tipo|nombre|props|metodos|eventos|integraciones|reglas"
    Ejemplo: "Servicio|UsuarioService|2|5|0|3|2"
    """
    global estimacion_actual

    if estimacion_actual is None:
        estimacion_actual = EstimacionProyecto("Proyecto Manual")

    try:
        partes = especificacion.split("|")
        if len(partes) < 7:
            return "Formato inválido. Usa: tipo|nombre|props|metodos|eventos|integraciones|reglas"

        comp = ComponenteEstimado(partes[0], partes[1])
        comp.propiedades = int(partes[2])
        comp.metodos = int(partes[3])
        comp.eventos = int(partes[4])
        comp.integraciones = int(partes[5])
        comp.reglas_negocio = int(partes[6])

        estimacion_actual.agregar_componente(comp)
        return f"Componente '{comp.nombre}' agregado correctamente"

    except Exception as e:
        return f"Error al agregar componente: {str(e)}"


def calcular_estimacion_completa(parametros: str) -> str:
    """
    Calcula la estimación completa aplicando ponderaciones y factores.
    Parámetros: "incertidumbre|acoplamiento|seniority"
    Ejemplo: "Media|Si|Mid"
    """
    global estimacion_actual

    if estimacion_actual is None or not estimacion_actual.componentes:
        return "No hay componentes para estimar. Primero extrae o agrega componentes."

    try:
        # Parsear parámetros
        partes = parametros.split("|")
        if len(partes) >= 3:
            estimacion_actual.incertidumbre = partes[0]
            estimacion_actual.acoplamiento = partes[1].lower() in ['si', 'sí', 'yes', 'true']
            estimacion_actual.seniority = partes[2]

        # Leer ponderaciones del Excel
        ponderaciones = pd.read_excel(EXCEL_TEMPLATE_PATH, sheet_name='Ponderaciones', header=0)

        # Calcular horas de cada componente
        for comp in estimacion_actual.componentes:
            comp.calcular_horas(ponderaciones)

        # Aplicar factores de ajuste
        totales = estimacion_actual.aplicar_factores_ajuste(ponderaciones)

        # Distribuir en fases - REDONDEO AL SIGUIENTE ENTERO
        total_desarrollo = totales['ajustada_seniority']
        estimacion_actual.horas_construccion = math.ceil(total_desarrollo * 0.60)
        estimacion_actual.horas_pruebas = math.ceil(total_desarrollo * 0.20)
        estimacion_actual.horas_analisis = math.ceil(total_desarrollo * 0.10)
        estimacion_actual.horas_documentacion = math.ceil(total_desarrollo * 0.05)
        estimacion_actual.horas_peer_review = math.ceil(total_desarrollo * 0.05)

        resultado = f"""
ESTIMACIÓN CALCULADA:
=====================
Componentes analizados: {len(estimacion_actual.componentes)}
Horas base: {totales['base']}h
Horas ajustadas (seniority {estimacion_actual.seniority}): {totales['ajustada_seniority']}h
+ Incertidumbre ({estimacion_actual.incertidumbre}): {totales['incertidumbre']}h
+ Acoplamiento: {totales['acoplamiento']}h
─────────────────────────────────
TOTAL FINAL: {totales['total_final']} horas

Distribución por fases:
- Análisis técnico: {estimacion_actual.horas_analisis}h
- Construcción: {estimacion_actual.horas_construccion}h
- Pruebas: {estimacion_actual.horas_pruebas}h
- Documentación: {estimacion_actual.horas_documentacion}h
- Peer Review: {estimacion_actual.horas_peer_review}h
"""
        return resultado

    except Exception as e:
        return f"Error al calcular estimación: {str(e)}"


def exportar_a_excel(nombre_archivo_salida: str) -> str:
    """
    Exporta la estimación al archivo Excel GESTIC.
    Args:
        nombre_archivo_salida: Nombre para el archivo de salida (ej: "estimacion_proyecto_x.xlsx")
    """
    global estimacion_actual

    if estimacion_actual is None:
        return "No hay estimación para exportar"

    try:
        # Guardar en carpeta OUTPUT
        ruta_salida = os.path.join(PROJECT_ROOT, "OUTPUT", nombre_archivo_salida)

        # Cargar el workbook
        wb = load_workbook(EXCEL_TEMPLATE_PATH)

        # ===== ESCRIBIR EN ESTIMADOR_RESUMEN =====
        ws_resumen = wb['Estimador_resumen']

        # Buscar la fila donde empiezan los datos (fila con "NOMBRE")
        fila_inicio = 4  # Aproximadamente fila 4 según la lectura anterior

        # Escribir valores en la columna "Horas final" (columna D)
        ws_resumen[f'D{fila_inicio+1}'] = estimacion_actual.horas_gestion  # Gestión
        ws_resumen[f'D{fila_inicio+2}'] = estimacion_actual.horas_analisis  # Análisis técnico
        ws_resumen[f'D{fila_inicio+3}'] = estimacion_actual.horas_construccion  # Construcción
        ws_resumen[f'D{fila_inicio+4}'] = estimacion_actual.horas_pruebas  # Pruebas
        ws_resumen[f'D{fila_inicio+5}'] = estimacion_actual.horas_documentacion  # Documentación
        ws_resumen[f'D{fila_inicio+6}'] = estimacion_actual.horas_peer_review  # Peer Review

        # Escribir Incertidumbre y Acoplamiento
        ws_resumen['C1'] = estimacion_actual.incertidumbre
        ws_resumen['C2'] = "Sí" if estimacion_actual.acoplamiento else "No"

        # ===== ESCRIBIR EN ESTIMADOR_DESGLOSADO =====
        ws_desglosado = wb['Estimador_desglosado']

        fila_actual = 6  # Empezar después de la fila de encabezados

        for comp in estimacion_actual.componentes:
            ws_desglosado[f'A{fila_actual}'] = comp.tipo
            ws_desglosado[f'B{fila_actual}'] = "SÍ" if comp.es_nuevo else "NO"
            ws_desglosado[f'C{fila_actual}'] = comp.nombre
            ws_desglosado[f'D{fila_actual}'] = comp.propiedades
            ws_desglosado[f'E{fila_actual}'] = comp.metodos
            ws_desglosado[f'F{fila_actual}'] = comp.eventos
            ws_desglosado[f'G{fila_actual}'] = comp.integraciones
            ws_desglosado[f'H{fila_actual}'] = comp.reglas_negocio
            ws_desglosado[f'I{fila_actual}'] = comp.horas_calculadas
            fila_actual += 1

        # Guardar
        wb.save(ruta_salida)

        return f"✅ Estimación exportada correctamente a:\n{ruta_salida}\n\nTotal componentes: {len(estimacion_actual.componentes)}\nTotal horas: {estimacion_actual.calcular_total_horas():.2f}h"

    except Exception as e:
        return f"Error al exportar a Excel: {str(e)}"


# ============================================================================
# CONFIGURACIÓN DEL AGENTE
# ============================================================================

def crear_agente_estimacion():
    """Crea el agente de estimación con todas sus herramientas"""

    # Inicializar LLM
    llm = Ollama(
        model="llama2",  # o mistral, llama3.2
        temperature=0.3  # Baja temperatura para ser más preciso
    )

    # Definir herramientas
    herramientas = [
        Tool(
            name="LeerPDF",
            func=leer_pdf,
            description="Lee un archivo PDF de diseño técnico. Input: ruta completa al archivo PDF"
        ),
        Tool(
            name="LeerWord",
            func=leer_word,
            description="Lee un archivo Word (.docx) de diseño técnico. Input: ruta completa al archivo"
        ),
        Tool(
            name="ExtraerComponentes",
            func=extraer_componentes_tecnicos,
            description="Analiza un diseño técnico y extrae componentes automáticamente. Input: texto del diseño"
        ),
        Tool(
            name="AgregarComponente",
            func=agregar_componente_manual,
            description="Agrega un componente manualmente. Input: 'tipo|nombre|props|metodos|eventos|integraciones|reglas'"
        ),
        Tool(
            name="CalcularEstimacion",
            func=calcular_estimacion_completa,
            description="Calcula estimación completa. Input: 'incertidumbre|acoplamiento|seniority' (ej: 'Media|Si|Mid')"
        ),
        Tool(
            name="ExportarExcel",
            func=exportar_a_excel,
            description="Exporta la estimación a Excel GESTIC. Input: nombre del archivo de salida (ej: 'proyecto_x.xlsx')"
        )
    ]

    # Prompt para el agente (integrado con el prompt de estimación)
    template = f"""Eres un Agente de Estimación de Desarrollo de Software experto.

{PROMPT_ESTIMACION[:1000]}  # Incluir parte del prompt original

Tienes las siguientes herramientas:
{{tools}}

Nombres: {{tool_names}}

PROCESO RECOMENDADO:
1. Si el usuario proporciona un archivo, usa LeerPDF o LeerWord
2. Analiza el contenido con ExtraerComponentes
3. Agrega componentes manualmente si es necesario con AgregarComponente
4. Calcula la estimación con CalcularEstimacion
5. Exporta a Excel con ExportarExcel

Formato de respuesta:
Pregunta: {{input}}
Pensamiento: [qué necesito hacer]
Acción: [herramienta a usar]
Entrada de Acción: [input para la herramienta]
Observación: [resultado]
... (repetir si es necesario)
Pensamiento: Ahora tengo la respuesta final
Respuesta Final: [respuesta al usuario]

Historial: {{chat_history}}
Pregunta: {{input}}
{{agent_scratchpad}}"""

    prompt = PromptTemplate(
        template=template,
        input_variables=["input", "chat_history", "agent_scratchpad", "tools", "tool_names"]
    )

    # Memoria
    memoria = ConversationBufferMemory(
        memory_key="chat_history",
        return_messages=True
    )

    # Crear agente
    agente = create_react_agent(llm, herramientas, prompt)

    # Ejecutor
    agente_executor = AgentExecutor(
        agent=agente,
        tools=herramientas,
        memory=memoria,
        verbose=True,
        handle_parsing_errors=True,
        max_iterations=10
    )

    return agente_executor


# ============================================================================
# MAIN
# ============================================================================

def main():
    print("="*80)
    print("🤖 AGENTE DE ESTIMACIÓN DE DESARROLLO - LANGCHAIN + OLLAMA")
    print("="*80)
    print("\nCapacidades:")
    print("  📄 Leer diseños técnicos (PDF, Word)")
    print("  🧠 Analizar componentes automáticamente")
    print("  📊 Calcular estimaciones con ponderaciones GESTIC")
    print("  📤 Exportar a plantilla Excel GESTIC")
    print("\nEscribe 'salir' para terminar\n")
    print("="*80 + "\n")

    try:
        agente = crear_agente_estimacion()
        print("✅ Agente inicializado\n")
    except Exception as e:
        print(f"❌ Error: {e}")
        return

    while True:
        try:
            entrada = input("👤 Tú: ").strip()

            if entrada.lower() in ['salir', 'exit', 'quit']:
                print("\n👋 ¡Hasta luego!")
                break

            if not entrada:
                continue

            print("\n🤖 Agente:")
            respuesta = agente.invoke({"input": entrada})
            print("\n" + "="*80)
            print(f"📤 Respuesta: {respuesta['output']}")
            print("="*80 + "\n")

        except KeyboardInterrupt:
            print("\n\n👋 ¡Hasta luego!")
            break
        except Exception as e:
            print(f"\n❌ Error: {e}\n")


if __name__ == "__main__":
    main()
